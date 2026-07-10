from __future__ import annotations

from pathlib import Path

import fitz
from docx import Document
from PIL import Image

from edumind.ocr.core.base_extractor import ExtractionResult
from edumind.ocr.core.pipeline import DataIngestionPipeline
from edumind.ocr.extractors._image_backends import ImageOCRBackends, OCRRunResult


def _create_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        (
            "This is a generated study note PDF used for OCR pipeline integration testing. "
            "It contains enough text to stay on the normal PyMuPDF path."
        ),
    )
    doc.save(path)
    doc.close()


def _create_scanned_pdf(path: Path) -> None:
    doc = fitz.open()
    doc.new_page()
    doc.save(path)
    doc.close()


def _create_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Study Notes", level=1)
    document.add_paragraph("Linear algebra focuses on vectors, matrices, and transformations.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Topic"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Chapter"
    table.rows[1].cells[1].text = "1"
    document.save(path)


def _create_image(path: Path) -> None:
    Image.new("RGB", (120, 60), color="white").save(path)


class _StubImageExtractor:
    def extract(self, file_path: Path, **kwargs) -> ExtractionResult:
        return ExtractionResult(
            text="Name: Alice Example\n123\nImage notes with $x+y$",
            metadata={
                "extractor": "stub-image",
                "ocr_data": {
                    "text": ["Name", "Alice", "Example"],
                    "conf": [95.0, 94.0, 94.0],
                    "left": [0, 10, 40],
                    "top": [0, 0, 0],
                    "width": [10, 20, 20],
                    "height": [10, 10, 10],
                },
                "image_shape": [30, 60, 3],
            },
            format_type="image",
            file_path=str(file_path),
            success=True,
        )


def test_pipeline_processes_pdf_docx_and_image_fixtures(tmp_path: Path) -> None:
    pdf_path = tmp_path / "notes.pdf"
    docx_path = tmp_path / "notes.docx"
    image_path = tmp_path / "notes.png"
    _create_pdf(pdf_path)
    _create_docx(docx_path)
    _create_image(image_path)

    pipeline = DataIngestionPipeline()
    pipeline.extractors["image"] = _StubImageExtractor()

    pdf_result = pipeline.process_file(pdf_path)
    docx_result = pipeline.process_file(docx_path)
    image_result = pipeline.process_file(image_path)

    assert pdf_result.success is True
    assert "generated study note PDF" in pdf_result.text
    assert pdf_result.metadata["format_info"]["format_type"] == "pdf"

    assert docx_result.success is True
    assert "Study Notes" in docx_result.text
    assert docx_result.metadata["num_tables"] == 1

    assert image_result.success is True
    assert image_result.metadata["format_info"]["format_type"] == "image"
    assert image_result.metadata["math_expressions"]["inline"] == ["x+y"]
    assert "file_hash" in image_result.metadata


def test_pipeline_adds_structured_metadata_and_skips_hash_when_disabled(tmp_path: Path) -> None:
    image_path = tmp_path / "structured.png"
    _create_image(image_path)

    pipeline = DataIngestionPipeline()
    pipeline.extractors["image"] = _StubImageExtractor()
    result = pipeline.process_file(
        image_path,
        include_layout=True,
        include_form_fields=True,
        include_file_hash=False,
        profile=True,
    )

    assert result.metadata["structured_fields"]["name"]["value"] == "Alice Example"
    assert result.metadata["layout_blocks"]
    assert "file_hash" not in result.metadata
    assert "performance" in result.metadata


def test_scanned_pdf_second_run_hits_page_cache(tmp_path: Path, monkeypatch) -> None:
    scanned_pdf = tmp_path / "scanned.pdf"
    _create_scanned_pdf(scanned_pdf)

    def _fake_extract(self, image, **kwargs) -> OCRRunResult:
        return OCRRunResult(
            text="Recovered scanned PDF text",
            confidence=95.0,
            attempts=["standard (conf: 95.00)"],
            ocr_data=None,
        )

    monkeypatch.setattr(ImageOCRBackends, "extract", _fake_extract)

    pipeline = DataIngestionPipeline()
    first = pipeline.process_file(scanned_pdf, pdf_ocr_mode="force")
    second = pipeline.process_file(scanned_pdf, pdf_ocr_mode="force")

    assert first.metadata["cache"]["page_misses"] == 1
    assert second.metadata["cache"]["page_hits"] == 1
    assert "Recovered scanned PDF text" in second.text

"""Lightweight native DOCX extraction controls."""

from __future__ import annotations

import time

from ..contracts import ExtractedDocument, ExtractionRequest, SourceKind
from ..errors import ExtractionBackendError, MissingDependencyError
from .base import build_document


class DOCXExtractor:
    supported_kinds = frozenset({SourceKind.DOCX})

    def __init__(self, engine: str, revision: str = "unpinned") -> None:
        self.engine = engine
        self.name = engine
        self.revision = revision

    def extract(self, request: ExtractionRequest, kind: SourceKind) -> ExtractedDocument:
        if request.profile is None:
            raise ValueError("Resolved extraction profile is required")
        started = time.perf_counter()
        try:
            text = self._extract(request)
        except MissingDependencyError:
            raise
        except Exception as exc:
            raise ExtractionBackendError(
                f"DOCX extraction failed with {self.engine}", detail=str(exc)
            ) from exc
        return build_document(
            request,
            kind,
            request.profile,
            [text.strip()],
            metadata={"engine": self.engine},
            seconds=time.perf_counter() - started,
        )

    def _extract(self, request: ExtractionRequest) -> str:
        if self.engine == "python-docx":
            try:
                from docx import Document
            except ModuleNotFoundError as exc:
                raise MissingDependencyError(
                    "python-docx is required; install requirements/app.lock"
                ) from exc
            document = Document(str(request.source_path))
            blocks = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                blocks.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
            return "\n".join(blocks)
        if self.engine == "mammoth":
            try:
                import mammoth
            except ModuleNotFoundError as exc:
                raise MissingDependencyError("Mammoth is required for this DOCX candidate") from exc
            with request.source_path.open("rb") as handle:
                return str(mammoth.extract_raw_text(handle).value)
        raise ValueError(f"Unknown DOCX engine: {self.engine}")

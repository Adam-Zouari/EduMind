"""DOCX extraction using python-docx."""

from __future__ import annotations

import time
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument

from ..core.base_extractor import BaseExtractor, ExtractionResult


class DOCXExtractor(BaseExtractor):
    """Extract text and metadata from DOCX files."""

    def extract(self, file_path: Path, **kwargs: object) -> ExtractionResult:
        """Extract text from a DOCX file."""
        start_time = time.time()
        self.logger.info(f"Extracting DOCX: {file_path}")

        try:
            document = Document(str(file_path))
            paragraphs = self._extract_paragraphs(document)
            tables = self._extract_tables(document)
            return ExtractionResult(
                text=self._combine_text(paragraphs, tables),
                metadata=self._build_metadata(document),
                format_type="docx",
                file_path=str(file_path),
                extraction_time=time.time() - start_time,
                success=True,
            )
        except Exception as exc:
            self.logger.error(f"DOCX extraction failed: {exc}")
            return self._create_error_result(file_path, str(exc))

    def _extract_paragraphs(self, document: DocxDocument) -> list[str]:
        """Extract non-empty paragraph text."""
        return [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    def _extract_tables(self, document: DocxDocument) -> list[str]:
        """Extract table rows as plain text lines."""
        rows: list[str] = []
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    rows.append(row_text)
        return rows

    def _combine_text(self, paragraphs: list[str], tables: list[str]) -> str:
        """Combine paragraph and table text into the existing OCR payload shape."""
        text = "\n\n".join(paragraphs)
        if tables:
            text += "\n\n--- TABLES ---\n\n" + "\n".join(tables)
        return text

    def _build_metadata(self, document: DocxDocument) -> dict[str, object]:
        """Build the normalized DOCX metadata payload."""
        core_props = document.core_properties
        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "num_paragraphs": len(document.paragraphs),
            "num_tables": len(document.tables),
            "extractor": "python-docx",
        }

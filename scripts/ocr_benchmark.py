"""Repeatable OCR benchmark runner for the local OCR package."""

from __future__ import annotations

import csv
import json
import math
from pathlib import Path
import struct
import subprocess
import time
import wave

from docx import Document
import fitz
from PIL import Image, ImageDraw

from edumind.common.paths import artifact_path
from edumind.ocr.core.pipeline import DataIngestionPipeline

BENCHMARK_DIR = artifact_path("ocr", "benchmarks")
CORPUS_DIR = artifact_path("ocr", "benchmarks", "corpus")
BENCHMARK_DIR.mkdir(parents=True, exist_ok=True)
CORPUS_DIR.mkdir(parents=True, exist_ok=True)


def main() -> None:
    pipeline = DataIngestionPipeline()
    corpus = ensure_corpus()
    timestamp = time.strftime("%Y%m%d_%H%M%S")

    scenarios = [
        {
            "name": "native_pdf_default",
            "type": "file",
            "target": corpus["native_pdf"],
            "kwargs": {"profile": True},
        },
        {
            "name": "scanned_pdf_native_only",
            "type": "file",
            "target": corpus["scanned_pdf"],
            "kwargs": {"profile": True, "pdf_ocr_mode": "off"},
        },
        {
            "name": "scanned_pdf_force_ocr_first",
            "type": "file",
            "target": corpus["scanned_pdf"],
            "kwargs": {"profile": True, "pdf_ocr_mode": "force"},
        },
        {
            "name": "scanned_pdf_force_ocr_second",
            "type": "file",
            "target": corpus["scanned_pdf"],
            "kwargs": {"profile": True, "pdf_ocr_mode": "force"},
        },
        {
            "name": "docx_default",
            "type": "file",
            "target": corpus["docx"],
            "kwargs": {"profile": True},
        },
        {
            "name": "clean_image_first",
            "type": "file",
            "target": corpus["clean_image"],
            "kwargs": {"profile": True},
        },
        {
            "name": "clean_image_second",
            "type": "file",
            "target": corpus["clean_image"],
            "kwargs": {"profile": True},
        },
        {
            "name": "noisy_image_default",
            "type": "file",
            "target": corpus["noisy_image"],
            "kwargs": {"profile": True},
        },
        {
            "name": "mixed_batch_threads",
            "type": "batch",
            "target": [
                corpus["native_pdf"],
                corpus["docx"],
                corpus["clean_image"],
                corpus["noisy_image"],
            ],
            "kwargs": {"profile": True, "batch_strategy": "threads"},
        },
        {
            "name": "mixed_batch_sequential",
            "type": "batch",
            "target": [
                corpus["native_pdf"],
                corpus["docx"],
                corpus["clean_image"],
                corpus["noisy_image"],
            ],
            "kwargs": {"profile": True, "batch_strategy": "sequential"},
        },
    ]

    if corpus["audio"] is not None:
        scenarios.append(
            {
                "name": "short_audio_default",
                "type": "file",
                "target": corpus["audio"],
                "kwargs": {"profile": True},
            }
        )
    if corpus["video"] is not None:
        scenarios.append(
            {
                "name": "short_video_default",
                "type": "file",
                "target": corpus["video"],
                "kwargs": {"profile": True},
            }
        )

    results = [run_scenario(pipeline, scenario) for scenario in scenarios]
    summary = build_summary(results)

    payload = {
        "generated_at": timestamp,
        "corpus": {key: str(value) if value is not None else None for key, value in corpus.items()},
        "results": results,
        "summary": summary,
    }

    json_path = BENCHMARK_DIR / f"benchmark_{timestamp}.json"
    csv_path = BENCHMARK_DIR / f"benchmark_{timestamp}.csv"
    json_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    write_csv(results, csv_path)

    print(f"Benchmark JSON written to {json_path}")
    print(f"Benchmark CSV written to {csv_path}")
    print(json.dumps(summary, indent=2))


def ensure_corpus() -> dict[str, Path | None]:
    """Create or reuse the local benchmark corpus."""
    native_pdf = CORPUS_DIR / "native_text.pdf"
    scanned_pdf = CORPUS_DIR / "scanned_text.pdf"
    docx_path = CORPUS_DIR / "study_notes.docx"
    clean_image = CORPUS_DIR / "clean_text.png"
    noisy_image = CORPUS_DIR / "noisy_text.png"
    audio_path = CORPUS_DIR / "short_tone.wav"
    video_path = CORPUS_DIR / "short_video.mp4"

    if not native_pdf.exists():
        create_native_pdf(native_pdf)
    if not clean_image.exists():
        create_clean_image(clean_image)
    if not noisy_image.exists():
        create_noisy_image(noisy_image)
    if not scanned_pdf.exists():
        create_scanned_pdf(scanned_pdf, clean_image)
    if not docx_path.exists():
        create_docx(docx_path)
    if not audio_path.exists():
        create_audio(audio_path)
    if not video_path.exists():
        if not create_video(video_path, audio_path):
            video_path = None

    return {
        "native_pdf": native_pdf,
        "scanned_pdf": scanned_pdf,
        "docx": docx_path,
        "clean_image": clean_image,
        "noisy_image": noisy_image,
        "audio": audio_path if audio_path.exists() else None,
        "video": video_path,
    }


def create_native_pdf(path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        (
            "This native PDF benchmark represents a study handout with headings, "
            "sentences, and enough searchable text to remain on the non-OCR path."
        ),
    )
    doc.save(path)
    doc.close()


def create_scanned_pdf(path: Path, image_path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_image(page.rect, filename=str(image_path))
    doc.save(path)
    doc.close()


def create_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Benchmark Study Notes", level=1)
    document.add_paragraph(
        "Differentiation studies change, while integration studies accumulation."
    )
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Topic"
    table.rows[0].cells[1].text = "Chapter"
    table.rows[1].cells[0].text = "Calculus"
    table.rows[1].cells[1].text = "2"
    document.save(path)


def create_clean_image(path: Path) -> None:
    image = Image.new("RGB", (800, 400), color="white")
    draw = ImageDraw.Draw(image)
    draw.text((40, 60), "Study Guide\nLinear Algebra and Calculus", fill="black")
    draw.text((40, 180), "Name: Alice Example", fill="black")
    draw.text((40, 240), "Equation: x^2 + y^2 = z^2", fill="black")
    image.save(path)


def create_noisy_image(path: Path) -> None:
    clean_path = CORPUS_DIR / "clean_text.png"
    if not clean_path.exists():
        create_clean_image(clean_path)

    image = Image.open(clean_path).convert("RGB")
    pixels = image.load()
    width, height = image.size
    for x in range(width):
        for y in range(height):
            if (x * y) % 17 == 0:
                r, g, b = pixels[x, y]
                pixels[x, y] = (
                    min(255, r + 30),
                    max(0, g - 20),
                    min(255, b + 10),
                )
    image.save(path)


def create_audio(path: Path) -> None:
    sample_rate = 16000
    duration_seconds = 1.0
    frequency = 440.0
    amplitude = 12000

    with wave.open(str(path), "w") as wav_file:
        wav_file.setnchannels(1)
        wav_file.setsampwidth(2)
        wav_file.setframerate(sample_rate)

        for index in range(int(sample_rate * duration_seconds)):
            sample = int(amplitude * math.sin(2 * math.pi * frequency * index / sample_rate))
            wav_file.writeframes(struct.pack("<h", sample))


def create_video(path: Path, audio_path: Path) -> bool:
    """Create a short local MP4 if ffmpeg is available."""
    command = [
        "ffmpeg",
        "-y",
        "-f",
        "lavfi",
        "-i",
        "color=c=black:s=320x240:d=1",
        "-i",
        str(audio_path),
        "-shortest",
        "-pix_fmt",
        "yuv420p",
        str(path),
    ]
    try:
        subprocess.run(command, check=True, capture_output=True)
        return True
    except Exception:
        return False


def run_scenario(pipeline: DataIngestionPipeline, scenario: dict[str, object]) -> dict[str, object]:
    """Run one benchmark scenario and capture summary fields."""
    start = time.perf_counter()
    name = str(scenario["name"])
    scenario_type = str(scenario["type"])
    kwargs = dict(scenario["kwargs"])

    try:
        if scenario_type == "file":
            result = pipeline.process_file(Path(scenario["target"]), **kwargs)
            elapsed = time.perf_counter() - start
            return {
                "name": name,
                "type": scenario_type,
                "target": str(scenario["target"]),
                "elapsed_seconds": elapsed,
                "success": result.success,
                "text_length": len(result.text),
                "format_type": result.format_type,
                "cache": result.metadata.get("cache"),
                "performance": result.metadata.get("performance"),
                "error": result.error,
            }

        results = pipeline.process_batch([Path(item) for item in scenario["target"]], **kwargs)
        elapsed = time.perf_counter() - start
        return {
            "name": name,
            "type": scenario_type,
            "target": [str(item) for item in scenario["target"]],
            "elapsed_seconds": elapsed,
            "success": all(result.success for result in results),
            "success_count": sum(1 for result in results if result.success),
            "result_count": len(results),
            "text_lengths": [len(result.text) for result in results],
            "formats": [result.format_type for result in results],
            "errors": [result.error for result in results if result.error],
        }
    except Exception as exc:
        return {
            "name": name,
            "type": scenario_type,
            "target": scenario["target"],
            "elapsed_seconds": time.perf_counter() - start,
            "success": False,
            "error": str(exc),
        }


def build_summary(results: list[dict[str, object]]) -> dict[str, object]:
    """Build comparison-friendly benchmark summary fields."""
    by_name = {result["name"]: result for result in results}
    clean_first = float(by_name.get("clean_image_first", {}).get("elapsed_seconds", 0.0) or 0.0)
    clean_second = float(by_name.get("clean_image_second", {}).get("elapsed_seconds", 0.0) or 0.0)
    scanned_first = float(by_name.get("scanned_pdf_force_ocr_first", {}).get("elapsed_seconds", 0.0) or 0.0)
    scanned_second = float(by_name.get("scanned_pdf_force_ocr_second", {}).get("elapsed_seconds", 0.0) or 0.0)
    threads = float(by_name.get("mixed_batch_threads", {}).get("elapsed_seconds", 0.0) or 0.0)
    sequential = float(by_name.get("mixed_batch_sequential", {}).get("elapsed_seconds", 0.0) or 0.0)

    return {
        "cache_rerun_improvement_clean_image_pct": percent_improvement(clean_first, clean_second),
        "cache_rerun_improvement_scanned_pdf_pct": percent_improvement(scanned_first, scanned_second),
        "mixed_batch_threads_vs_sequential_pct": percent_improvement(sequential, threads),
        "native_vs_scanned_pdf": {
            "native_pdf_default_seconds": by_name.get("native_pdf_default", {}).get("elapsed_seconds"),
            "scanned_pdf_native_only_seconds": by_name.get("scanned_pdf_native_only", {}).get("elapsed_seconds"),
            "scanned_pdf_force_ocr_seconds": by_name.get("scanned_pdf_force_ocr_first", {}).get("elapsed_seconds"),
        },
    }


def percent_improvement(baseline: float, comparison: float) -> float | None:
    """Calculate percent improvement when lower time is better."""
    if baseline <= 0:
        return None
    return round(((baseline - comparison) / baseline) * 100, 2)


def write_csv(results: list[dict[str, object]], path: Path) -> None:
    """Write a flat CSV summary of benchmark scenarios."""
    fieldnames = [
        "name",
        "type",
        "target",
        "elapsed_seconds",
        "success",
        "text_length",
        "format_type",
        "result_count",
        "success_count",
        "error",
    ]
    with path.open("w", encoding="utf-8", newline="") as csv_file:
        writer = csv.DictWriter(csv_file, fieldnames=fieldnames)
        writer.writeheader()
        for result in results:
            writer.writerow(
                {
                    "name": result.get("name"),
                    "type": result.get("type"),
                    "target": result.get("target"),
                    "elapsed_seconds": result.get("elapsed_seconds"),
                    "success": result.get("success"),
                    "text_length": result.get("text_length"),
                    "format_type": result.get("format_type"),
                    "result_count": result.get("result_count"),
                    "success_count": result.get("success_count"),
                    "error": result.get("error"),
                }
            )


if __name__ == "__main__":
    main()

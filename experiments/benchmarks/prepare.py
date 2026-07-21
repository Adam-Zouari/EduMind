"""Explicit, checksum-validating dataset preparation; never runs at import time."""

from __future__ import annotations

import argparse
import json
import random
import subprocess
from collections import Counter
import sys
from collections.abc import Mapping, Sequence
from importlib.metadata import version
from pathlib import Path
from typing import Any

import requests

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from edumind.common.artifacts import atomic_write_json, atomic_write_text, sha256_file
from edumind.extraction.normalization import normalize_text

from experiments.benchmarks.common.datasets import (
    assert_no_split_leakage,
    load_manifest,
    manifest_content_checksum,
)

QASPER_DATASET = "allenai/qasper"
QASPER_REVISION = "3065362e337ded696bbb0171b073c73e513c9410"
HUGGINGFACE_MODELS = (
    "sentence-transformers/all-MiniLM-L6-v2",
    "google/embeddinggemma-300m",
    "infgrad/Jasper-Token-Compression-600M",
    "Qwen/Qwen3-Embedding-0.6B",
    "nvidia/Nemotron-3-Embed-1B-BF16",
    "Qwen/Qwen3-Embedding-4B",
    "nvidia/Nemotron-3-Embed-8B-BF16",
    "cross-encoder/ms-marco-MiniLM-L-6-v2",
    "BAAI/bge-reranker-v2-m3",
    "Qwen/Qwen3-Reranker-0.6B",
    "Qwen/Qwen3-Reranker-4B",
    "vectara/hallucination_evaluation_model",
)
PINNED_HUGGINGFACE_REVISIONS = {
    "sentence-transformers/all-MiniLM-L6-v2": "c9745ed1d9f207416be6d2e6f8de32d1f16199bf",
}
OLLAMA_MODELS = (
    "qwen3:1.7b",
    "qwen3.5:4b-q4_K_M",
    "qwen3.5:9b-q4_K_M",
    "gemma4:12b-it-q4_K_M",
    "ministral-3:8b-instruct-2512-q4_K_M",
    "gpt-oss:20b",
)
FASTER_WHISPER_MODELS = {
    "faster-whisper-tiny-int8": "Systran/faster-whisper-tiny.en",
    "faster-whisper-base-int8": "Systran/faster-whisper-base.en",
    "faster-whisper-small-int8": "Systran/faster-whisper-small.en",
    "faster-whisper-small-float16": "Systran/faster-whisper-small.en",
    "faster-whisper-turbo-int8": "mobiuslabsgmbh/faster-whisper-large-v3-turbo",
}

EXTRACTION_HUGGINGFACE_MODELS = {
    "distil-whisper-large-v3.5": "distil-whisper/distil-large-v3.5",
    "parakeet-tdt-0.6b-v3": "nvidia/parakeet-tdt-0.6b-v3",
    "canary-qwen-2.5b": "nvidia/canary-qwen-2.5b",
    "glm-ocr": "zai-org/GLM-OCR",
    "mineru-2.5-pro": "opendatalab/MinerU2.5-Pro-2605-1.2B",
    "olmocr-2-7b": "allenai/olmOCR-2-7B-1025",
}

VECTOR_IMAGES = {
    "chroma": "chromadb/chroma:1.5.9",
    "qdrant": "qdrant/qdrant:v1.17.0",
    "weaviate": "cr.weaviate.io/semitechnologies/weaviate:1.38.2",
    "pgvector": "pgvector/pgvector:0.8.2-pg17-bookworm",
    "inspector": "alpine:3.21",
}


def prepare_app_models(root: Path) -> list[Path]:
    """Download only the provisional application's MiniLM, base ASR, and Qwen model."""
    huggingface_lock = root / "data/benchmarks/models/huggingface.json"
    prepare_huggingface_models(
        huggingface_lock,
        ["sentence-transformers/all-MiniLM-L6-v2"],
    )
    extraction_lock = root / "data/benchmarks/models/extraction.json"
    prepare_extraction_models(
        extraction_lock,
        root / "data/benchmarks/downloads/models",
        ["faster-whisper-base-int8"],
    )
    ollama_lock = root / "data/benchmarks/models/ollama.json"
    prepare_ollama_models(ollama_lock, ["qwen3:1.7b"])
    return [huggingface_lock, extraction_lock, ollama_lock]


def _merge_model_lock(path: Path, source: str, updates: Mapping[str, object]) -> None:
    existing: dict[str, object] = {}
    if path.is_file():
        payload = json.loads(path.read_text(encoding="utf-8"))
        models = payload.get("models") if isinstance(payload, Mapping) else None
        if not isinstance(models, Mapping):
            raise ValueError(f"Existing model lock is malformed: {path}")
        existing.update(models)
    existing.update(updates)
    atomic_write_json(
        path,
        {"schema_version": 1, "source": source, "models": existing},
    )


def prepare_qasper(output_directory: Path, *, seed: int = 42) -> list[Path]:
    """Download pinned QASPER and create paper-isolated dev/validation/locked manifests."""
    try:
        from datasets import load_dataset
    except ModuleNotFoundError as exc:
        raise RuntimeError("datasets is required; install requirements/benchmarks.lock") from exc
    dataset = load_dataset(QASPER_DATASET, "qasper", revision=QASPER_REVISION)
    plans = (
        ("dev", "train", 100),
        ("validation", "validation", 40),
        ("locked-test", "test", 40),
    )
    output_directory.mkdir(parents=True, exist_ok=True)
    outputs: list[Path] = []
    selected_ids: set[str] = set()
    for manifest_split, source_split, count in plans:
        papers = [dict(row) for row in dataset[source_split]]
        selected = _stratified_papers(papers, count, seed)
        paper_ids = {str(paper["id"]) for paper in selected}
        if selected_ids & paper_ids:
            raise ValueError("QASPER paper leakage detected across prepared splits")
        selected_ids.update(paper_ids)
        samples = _qasper_samples(selected)
        payload = {
            "name": f"qasper-{manifest_split}",
            "version": "1.0.0",
            "task": "rag",
            "split": manifest_split,
            "source": QASPER_DATASET,
            "license": "CC-BY-4.0",
            "revision": QASPER_REVISION,
            "checksum": manifest_content_checksum(samples),
            "preprocessing_version": "qasper-normalized-text-v1",
            "split_seed": seed,
            "selected_ids": sorted(paper_ids),
            "samples": samples,
        }
        output = output_directory / f"qasper-{manifest_split}.json"
        atomic_write_json(output, payload)
        outputs.append(output)
    assert_no_split_leakage([load_manifest(path) for path in outputs])
    return outputs


def prepare_rag_selection_manifest(
    qasper_path: Path, structured_path: Path, output_path: Path
) -> Path:
    """Combine one QASPER split with verified table/formula/mixed RAG samples."""
    qasper = load_manifest(qasper_path)
    structured = load_manifest(structured_path)
    if qasper.split != structured.split:
        raise ValueError(
            f"RAG source splits differ: {qasper.split!r} versus {structured.split!r}"
        )
    structured_questions = [
        sample for sample in structured.samples if sample.get("kind") == "question"
    ]
    required_types = {"table", "formula", "mixed"}
    invalid_types = sorted(
        {
            str(sample.get("evidence_type", ""))
            for sample in structured_questions
            if str(sample.get("evidence_type", "")) not in required_types | {"text"}
        }
    )
    if invalid_types:
        raise ValueError(
            "Structured RAG questions use unsupported evidence_type values: "
            + ", ".join(invalid_types)
        )
    evidence_counts = Counter(
        str(sample.get("evidence_type"))
        for sample in structured_questions
        if sample.get("answerable") and sample.get("evidence")
    )
    insufficient = {
        evidence_type: evidence_counts[evidence_type]
        for evidence_type in required_types
        if evidence_counts[evidence_type] < 10
    }
    if insufficient:
        raise ValueError(
            "Structured RAG manifests require at least 10 answerable questions with "
            "verified spans for each structural evidence type; received "
            + ", ".join(f"{name}={count}" for name, count in sorted(insufficient.items()))
        )
    combined = [*qasper.samples, *structured.samples]
    identifiers = [str(sample.get("id", "")) for sample in combined]
    duplicates = sorted(identifier for identifier, count in Counter(identifiers).items() if count > 1)
    if duplicates:
        raise ValueError(
            "RAG manifests contain duplicate IDs; namespace the structured corpus: "
            + ", ".join(duplicates[:10])
        )
    payload = {
        "name": f"edumind-rag-selection-{qasper.split}",
        "version": "2.0.0",
        "task": "rag",
        "split": qasper.split,
        "source": f"{qasper.source}+{structured.source}",
        "license": f"{qasper.license};{structured.license}",
        "revision": f"{qasper.revision}+{structured.revision}",
        "checksum": manifest_content_checksum(combined),
        "preprocessing_version": "structured-rag-markdown-v1",
        "split_seed": 42,
        "samples": combined,
    }
    atomic_write_json(output_path, payload)
    # Re-read to enforce the normal checksum/evidence-offset contract.
    resolved = load_manifest(output_path)
    assert_no_split_leakage([resolved])
    return output_path


def prepare_public_assets(plan_path: Path, output_directory: Path) -> list[Path]:
    """Download a licensed asset plan and reject every checksum/license mismatch."""
    import json

    payload = json.loads(plan_path.read_text(encoding="utf-8"))
    assets = payload.get("assets", [])
    if not isinstance(assets, list) or not assets:
        raise ValueError("Asset plan must contain a non-empty assets list")
    output_directory.mkdir(parents=True, exist_ok=True)
    downloaded: list[Path] = []
    for asset in assets:
        if not isinstance(asset, Mapping):
            raise ValueError("Every asset entry must be an object")
        url = str(asset.get("url", ""))
        filename = Path(str(asset.get("filename", ""))).name
        expected = str(asset.get("sha256", "")).casefold()
        license_name = str(asset.get("license", ""))
        if (
            not url.startswith("https://")
            or not filename
            or len(expected) != 64
            or not license_name
        ):
            raise ValueError("Asset entries require HTTPS URL, filename, SHA-256, and license")
        destination = output_directory / filename
        temporary = destination.with_suffix(destination.suffix + ".partial")
        try:
            with requests.get(url, stream=True, timeout=120) as response:
                response.raise_for_status()
                with temporary.open("wb") as handle:
                    for block in response.iter_content(chunk_size=1024 * 1024):
                        if block:
                            handle.write(block)
            if sha256_file(temporary) != expected:
                raise ValueError(f"Checksum mismatch for {filename}")
            temporary.replace(destination)
        finally:
            temporary.unlink(missing_ok=True)
        downloaded.append(destination)
    return downloaded


def prepare_huggingface_models(
    output_path: Path, selected: Sequence[str] | None = None
) -> Path:
    """Resolve immutable revisions, download sequentially, and write a model lock."""
    try:
        from huggingface_hub import HfApi, snapshot_download
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "huggingface-hub is required; install requirements/benchmarks.lock"
        ) from exc
    api = HfApi()
    models = _selected(HUGGINGFACE_MODELS, selected, "Hugging Face model")
    for model in models:
        revision = PINNED_HUGGINGFACE_REVISIONS.get(model) or str(api.model_info(model).sha)
        snapshot_download(repo_id=model, revision=revision)
        # Persist each completed candidate. An interrupted later download does not
        # invalidate already prepared models, and snapshot_download resumes partials.
        _merge_model_lock(output_path, "huggingface", {model: revision})
    return output_path


def prepare_extraction_models(
    output_path: Path,
    cache_directory: Path,
    selected: Sequence[str] | None = None,
) -> Path:
    """Prepare selected extraction candidates and persist every completed item."""
    try:
        from huggingface_hub import HfApi, snapshot_download
    except ModuleNotFoundError as exc:
        raise RuntimeError("huggingface-hub is required for model preparation") from exc

    cache_directory = cache_directory.expanduser().resolve()
    cache_directory.mkdir(parents=True, exist_ok=True)
    available = (
        *FASTER_WHISPER_MODELS,
        *EXTRACTION_HUGGINGFACE_MODELS,
        "openai-whisper-small-en",
        "paddleocr-v5-mobile",
        "paddleocr-v5-server",
        "pp-structure-v3",
        "paddleocr-vl-1.6",
        "docling",
    )
    chosen = set(_selected(available, selected, "extraction candidate"))
    api = HfApi()
    downloaded_repositories: dict[str, tuple[str, Path]] = {}
    extraction_snapshots: dict[str, tuple[str, Path]] = {}

    for candidate, repository in FASTER_WHISPER_MODELS.items():
        if candidate not in chosen:
            continue
        if repository not in downloaded_repositories:
            revision = str(api.model_info(repository).sha)
            local_directory = cache_directory / repository.replace("/", "--")
            snapshot_download(repo_id=repository, revision=revision, local_dir=local_directory)
            downloaded_repositories[repository] = (revision, local_directory)
        revision, local_directory = downloaded_repositories[repository]
        _merge_model_lock(
            output_path,
            "explicit-local-preparation",
            {
                candidate: {
                    "provider": "huggingface",
                    "model": repository,
                    "revision": revision,
                    "model_path": str(local_directory),
                }
            },
        )

    for candidate, repository in EXTRACTION_HUGGINGFACE_MODELS.items():
        if candidate not in chosen:
            continue
        revision = str(api.model_info(repository).sha)
        local_directory = cache_directory / repository.replace("/", "--")
        snapshot_download(repo_id=repository, revision=revision, local_dir=local_directory)
        extraction_snapshots[candidate] = (revision, local_directory)
        _merge_model_lock(
            output_path,
            "explicit-local-preparation",
            {
                candidate: {
                    "provider": "huggingface",
                    "model": repository,
                    "revision": revision,
                    "model_path": str(local_directory),
                }
            },
        )

    if "mineru-2.5-pro" in chosen:
        _, model_directory = extraction_snapshots["mineru-2.5-pro"]
        mineru_config = cache_directory / "mineru-2.5-pro.json"
        atomic_write_json(
            mineru_config,
            {
                "latex-delimiter-config": {
                    "display": {"left": "$$", "right": "$$"},
                    "inline": {"left": "$", "right": "$"},
                },
                "llm-aided-config": {"title_aided": {"enable": False}},
                "models-dir": {"pipeline": "", "vlm": str(model_directory)},
                "model-source": "local",
                "config_version": "1.3.2",
            },
        )
        revision, _ = extraction_snapshots["mineru-2.5-pro"]
        _merge_model_lock(
            output_path,
            "explicit-local-preparation",
            {
                "mineru-2.5-pro": {
                    "provider": "huggingface",
                    "model": EXTRACTION_HUGGINGFACE_MODELS["mineru-2.5-pro"],
                    "revision": revision,
                    "model_path": str(model_directory),
                    "mineru_config_path": str(mineru_config),
                }
            },
        )

    if "openai-whisper-small-en" in chosen:
        try:
            import whisper
        except ModuleNotFoundError as exc:
            raise RuntimeError("openai-whisper is required for this candidate") from exc
        whisper_directory = cache_directory / "openai-whisper"
        whisper_directory.mkdir(parents=True, exist_ok=True)
        whisper_model = whisper.load_model(
            "small.en", download_root=str(whisper_directory), device="cpu"
        )
        del whisper_model
        whisper_path = whisper_directory / "small.en.pt"
        if not whisper_path.is_file():
            raise RuntimeError("OpenAI Whisper did not create the expected small.en weight file")
        _merge_model_lock(
            output_path,
            "explicit-local-preparation",
            {
                "openai-whisper-small-en": {
                    "provider": "openai-whisper",
                    "model": "small.en",
                    "revision": sha256_file(whisper_path),
                    "model_path": str(whisper_path),
                }
            },
        )

    paddle_cache = Path.home() / ".paddlex" / "official_models"
    for candidate, detector, recognizer in (
        ("paddleocr-v5-mobile", "PP-OCRv5_mobile_det", "en_PP-OCRv5_mobile_rec"),
        ("paddleocr-v5-server", "PP-OCRv5_server_det", "PP-OCRv5_server_rec"),
    ):
        if candidate not in chosen:
            continue
        try:
            from paddleocr import PaddleOCR
        except ModuleNotFoundError as exc:
            raise RuntimeError("PaddleOCR is required for this candidate") from exc
        PaddleOCR(
            lang="en",
            text_detection_model_name=detector,
            text_recognition_model_name=recognizer,
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=False,
        )
        detection_path, recognition_path = paddle_cache / detector, paddle_cache / recognizer
        if not detection_path.is_dir() or not recognition_path.is_dir():
            raise RuntimeError(f"PaddleOCR did not prepare local model directories for {candidate}")
        _merge_model_lock(
            output_path,
            "explicit-local-preparation",
            {
                candidate: {
                    "provider": "paddleocr",
                    "model": f"{detector}+{recognizer}",
                    "revision": version("paddleocr"),
                    "text_detection_model_dir": str(detection_path),
                    "text_recognition_model_dir": str(recognition_path),
                }
            },
        )

    if {"pp-structure-v3", "paddleocr-vl-1.6"} & chosen:
        if "pp-structure-v3" in chosen:
            try:
                from paddleocr import PPStructureV3
            except (ImportError, ModuleNotFoundError) as exc:
                raise RuntimeError("PP-StructureV3 is unavailable in PaddleOCR") from exc
            PPStructureV3(device="cpu")
        if "paddleocr-vl-1.6" in chosen:
            try:
                from paddlex import create_pipeline
            except (ImportError, ModuleNotFoundError) as exc:
                raise RuntimeError("PaddleX OCR extras are required for PaddleOCR-VL") from exc
            create_pipeline(pipeline="PaddleOCR-VL-1.6", device="cpu")
        paddle_root = Path.home() / ".paddlex"
        if not paddle_root.is_dir() or not any(paddle_root.rglob("*")):
            raise RuntimeError("Paddle document pipelines did not populate the local cache")
        if "pp-structure-v3" in chosen:
            _merge_model_lock(
                output_path,
                "explicit-local-preparation",
                {
                    "pp-structure-v3": {
                        "provider": "paddlex",
                        "model": "PP-StructureV3",
                        "revision": version("paddleocr"),
                        "paddle_cache_dir": str(paddle_root),
                    }
                },
            )
        if "paddleocr-vl-1.6" in chosen:
            _merge_model_lock(
                output_path,
                "explicit-local-preparation",
                {
                    "paddleocr-vl-1.6": {
                        "provider": "paddlex",
                        "model": "PaddleOCR-VL-1.6",
                        "revision": f"paddleocr-{version('paddleocr')}+paddlex-{version('paddlex')}",
                        "paddle_cache_dir": str(paddle_root),
                    }
                },
            )

    if "docling" in chosen:
        docling_directory = cache_directory / "docling"
        subprocess.run(
            ["docling-tools", "models", "download", "--output-dir", str(docling_directory)],
            check=True,
        )
        if not docling_directory.is_dir() or not any(docling_directory.rglob("*")):
            raise RuntimeError("Docling model preparation produced an empty artifact directory")
        _merge_model_lock(
            output_path,
            "explicit-local-preparation",
            {
                "docling": {
                    "provider": "docling",
                    "model": "default-local-pipeline",
                    "revision": version("docling"),
                    "docling_artifacts_dir": str(docling_directory),
                }
            },
        )
    return output_path


def prepare_ollama_models(output_path: Path, selected: Sequence[str] | None = None) -> Path:
    """Pull the documented Ollama candidates sequentially and lock installed digests."""
    _prepare_tiktoken()
    models = _selected(OLLAMA_MODELS, selected, "Ollama model")
    for model in models:
        subprocess.run(["ollama", "pull", model], check=True)
    response = requests.get("http://127.0.0.1:11434/api/tags", timeout=30)
    response.raise_for_status()
    payload = response.json()
    rows = payload.get("models", []) if isinstance(payload, Mapping) else []
    installed = {
        str(row.get("name")): str(row.get("digest"))
        for row in rows
        if isinstance(row, Mapping) and row.get("name") and row.get("digest")
    }
    missing = [model for model in models if model not in installed]
    if missing:
        raise RuntimeError(f"Ollama did not report pulled models: {', '.join(missing)}")
    _merge_model_lock(
        output_path,
        "ollama-local",
        {model: installed[model] for model in models},
    )
    return output_path


def _selected(
    available: Sequence[str], selected: Sequence[str] | None, label: str
) -> tuple[str, ...]:
    if not selected:
        return tuple(available)
    unknown = sorted(set(selected) - set(available))
    if unknown:
        raise ValueError(f"Unknown {label}(s): {', '.join(unknown)}")
    return tuple(dict.fromkeys(selected))


def _prepare_tiktoken() -> None:
    """Resolve the benchmark token counter during explicit preparation, never a run."""
    try:
        import tiktoken
    except ModuleNotFoundError as exc:
        raise RuntimeError("tiktoken is required; install the application lock") from exc
    tiktoken.get_encoding("cl100k_base").encode("EduMind preparation check")


def prepare_vectordb(output_path: Path) -> Path:
    """Pull four pinned tags, resolve immutable digests, and write Compose overrides."""
    resolved: dict[str, str] = {}
    for name, image in VECTOR_IMAGES.items():
        subprocess.run(["docker", "pull", image], check=True)
        process = subprocess.run(
            ["docker", "image", "inspect", image, "--format", "{{index .RepoDigests 0}}"],
            check=True,
            capture_output=True,
            text=True,
        )
        digest_image = process.stdout.strip()
        if "@sha256:" not in digest_image:
            raise RuntimeError(f"Docker did not report an immutable digest for {image}")
        resolved[name] = digest_image
    packages = {}
    for package in ("chromadb", "qdrant-client", "weaviate-client", "psycopg", "psycopg-pool"):
        packages[package] = version(package)
    atomic_write_json(
        output_path,
        {"schema_version": 1, "images": resolved, "clients": packages},
    )
    environment = output_path.parents[3] / "experiments/benchmarks/vectordb/.env"
    environment.parent.mkdir(parents=True, exist_ok=True)
    atomic_write_text(
        environment,
        "\n".join(
            f"{name.upper()}_IMAGE={image}" for name, image in resolved.items()
        )
        + "\n",
    )
    return output_path


def prepare_smoke_fixtures() -> Path:
    """Generate tiny valid multimodal files; no fake extractor is used by smoke runs."""
    from PIL import Image, ImageDraw
    import fitz
    from docx import Document

    root = Path(__file__).resolve().parents[2]
    fixtures = root / "data/benchmarks/fixtures/extraction"
    manifest_path = root / "data/benchmarks/extraction/smoke.json"
    payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    for sample in payload["samples"]:
        kind = sample["kind"]
        destination = root / sample["source_path"]
        destination.parent.mkdir(parents=True, exist_ok=True)
        text = sample["reference"]
        if kind == "image":
            image = Image.new("RGB", (1400, 220), "white")
            ImageDraw.Draw(image).text((40, 80), text, fill="black")
            image.save(destination)
        elif kind == "pdf":
            document = fitz.open()
            page = document.new_page(width=595, height=842)
            page.insert_text((72, 100), text, fontsize=16)
            document.save(destination)
            document.close()
        elif kind == "docx":
            document = Document()
            for index, line in enumerate(text.splitlines()):
                document.add_heading(line, level=1) if index == 0 else document.add_paragraph(line)
            document.save(destination)
        elif kind == "audio":
            _windows_speech(text, destination)
        elif kind == "video":
            temporary_image = destination.with_suffix(".png")
            temporary_audio = destination.with_suffix(".wav")
            image = Image.new("RGB", (1280, 720), "white")
            ImageDraw.Draw(image).text((80, 320), text, fill="black")
            image.save(temporary_image)
            _windows_speech(text, temporary_audio)
            try:
                subprocess.run(
                    [
                        "ffmpeg", "-y", "-loop", "1", "-i", str(temporary_image),
                        "-i", str(temporary_audio), "-c:v", "libx264", "-tune", "stillimage",
                        "-c:a", "aac", "-pix_fmt", "yuv420p", "-shortest", str(destination),
                    ],
                    check=True,
                    capture_output=True,
                )
            finally:
                temporary_image.unlink(missing_ok=True)
                temporary_audio.unlink(missing_ok=True)
        sample["asset_sha256"] = sha256_file(destination)
    payload["checksum"] = manifest_content_checksum(payload["samples"])
    atomic_write_json(manifest_path, payload)
    return manifest_path


def _windows_speech(text: str, destination: Path) -> None:
    escaped_text = text.replace("'", "''")
    escaped_path = str(destination.resolve()).replace("'", "''")
    command = (
        "Add-Type -AssemblyName System.Speech; "
        "$s = New-Object System.Speech.Synthesis.SpeechSynthesizer; "
        f"$s.SetOutputToWaveFile('{escaped_path}'); $s.Speak('{escaped_text}'); $s.Dispose()"
    )
    subprocess.run(["powershell", "-NoProfile", "-Command", command], check=True)


def main() -> int:
    parser = argparse.ArgumentParser(description="Explicit EduMind benchmark preparation")
    parser.add_argument(
        "target",
        choices=("app-models", "qasper", "rag-selection", "assets", "huggingface-models", "ollama-models", "extraction-models", "vectordb", "smoke-fixtures"),
    )
    parser.add_argument("--plan", type=Path)
    parser.add_argument("--output", type=Path)
    parser.add_argument("--qasper-manifest", type=Path)
    parser.add_argument("--structured-manifest", type=Path)
    parser.add_argument(
        "--candidate",
        action="append",
        help="Prepare only this candidate; repeat the option for several candidates",
    )
    arguments = parser.parse_args()
    root = Path(__file__).resolve().parents[2]
    if arguments.target == "app-models":
        outputs = prepare_app_models(root)
    elif arguments.target == "qasper":
        outputs = prepare_qasper(arguments.output or root / "data/benchmarks/rag")
    elif arguments.target == "rag-selection":
        if not arguments.qasper_manifest or not arguments.structured_manifest or not arguments.output:
            parser.error(
                "rag-selection requires --qasper-manifest, --structured-manifest, and --output"
            )
        outputs = [
            prepare_rag_selection_manifest(
                arguments.qasper_manifest,
                arguments.structured_manifest,
                arguments.output,
            )
        ]
    elif arguments.target == "assets":
        if arguments.plan is None:
            parser.error("assets requires --plan")
        outputs = prepare_public_assets(arguments.plan, arguments.output or root / "data/benchmarks/raw")
    elif arguments.target == "huggingface-models":
        outputs = [
            prepare_huggingface_models(
                arguments.output or root / "data/benchmarks/models/huggingface.json",
                arguments.candidate,
            )
        ]
    elif arguments.target == "ollama-models":
        outputs = [
            prepare_ollama_models(
                arguments.output or root / "data/benchmarks/models/ollama.json",
                arguments.candidate,
            )
        ]
    elif arguments.target == "extraction-models":
        outputs = [
            prepare_extraction_models(
                arguments.output or root / "data/benchmarks/models/extraction.json",
                root / "data/benchmarks/downloads/models",
                arguments.candidate,
            )
        ]
    elif arguments.target == "vectordb":
        outputs = [prepare_vectordb(arguments.output or root / "data/benchmarks/models/vectordb.json")]
    else:
        outputs = [prepare_smoke_fixtures()]
    print(json.dumps([str(path) for path in outputs], indent=2))
    return 0


def load_model_lock(path: Path) -> dict[str, str]:
    import json

    if not path.is_file():
        raise RuntimeError(f"Missing model lock {path}; run the matching preparation command")
    payload = json.loads(path.read_text(encoding="utf-8"))
    models = payload.get("models", {})
    if not isinstance(models, Mapping) or not models:
        raise RuntimeError(f"Model lock is empty or malformed: {path}")
    return {str(name): str(revision) for name, revision in models.items()}


def load_extraction_model_lock(path: Path) -> dict[str, dict[str, str]]:
    import json

    if not path.is_file():
        raise RuntimeError(
            f"Missing extraction model lock {path}; run `python "
            "experiments/benchmarks/prepare.py extraction-models`"
        )
    payload = json.loads(path.read_text(encoding="utf-8"))
    raw_models = payload.get("models", {})
    if not isinstance(raw_models, Mapping) or not raw_models:
        raise RuntimeError(f"Extraction model lock is empty or malformed: {path}")
    result: dict[str, dict[str, str]] = {}
    for candidate, raw_entry in raw_models.items():
        if not isinstance(raw_entry, Mapping) or not raw_entry.get("revision"):
            raise RuntimeError(f"Malformed extraction model lock entry: {candidate}")
        entry = {str(key): str(value) for key, value in raw_entry.items()}
        for key, value in entry.items():
            if key.endswith("_path") or key.endswith("_dir"):
                if not Path(value).exists():
                    raise RuntimeError(f"Prepared model path no longer exists: {value}")
        result[str(candidate)] = entry
    return result


def _stratified_papers(
    papers: Sequence[Mapping[str, Any]], count: int, seed: int
) -> list[Mapping[str, Any]]:
    if len(papers) < count:
        raise ValueError(f"Requested {count} papers from a split containing {len(papers)}")
    buckets: dict[tuple[int, bool], list[Mapping[str, Any]]] = {}
    for paper in papers:
        qas = _records(paper.get("qas", []))
        bucket = (min(3, len(qas) // 4), any(_question_unanswerable(qa) for qa in qas))
        buckets.setdefault(bucket, []).append(paper)
    random_state = random.Random(seed)
    for values in buckets.values():
        values.sort(key=lambda item: str(item["id"]))
        random_state.shuffle(values)
    ordered: list[Mapping[str, Any]] = []
    keys = sorted(buckets)
    while len(ordered) < count:
        progressed = False
        for key in keys:
            if buckets[key]:
                ordered.append(buckets[key].pop())
                progressed = True
                if len(ordered) == count:
                    break
        if not progressed:  # pragma: no cover - guarded by the length check
            break
    return ordered


def _qasper_samples(papers: Sequence[Mapping[str, Any]]) -> list[dict[str, object]]:
    samples: list[dict[str, object]] = []
    for paper in papers:
        paper_id = str(paper["id"])
        text = _paper_text(paper)
        samples.append(
            {
                "id": paper_id,
                "kind": "document",
                "text": text,
                "representation": "markdown-sections",
            }
        )
        for qa in _records(paper.get("qas", [])):
            answers, evidence, answer_type, answerable = _answers_and_evidence(qa, text, paper_id)
            samples.append(
                {
                    "id": str(qa.get("question_id", "")),
                    "kind": "question",
                    "document_id": paper_id,
                    "question": str(qa.get("question", "")),
                    "answer": answers[0] if answers else "",
                    "accepted_answers": answers,
                    "answer_type": answer_type,
                    "answerable": answerable,
                    "evidence_type": "text",
                    "evidence": evidence,
                }
            )
    return samples


def _paper_text(paper: Mapping[str, Any]) -> str:
    sections = _records(paper.get("full_text", []))
    title = str(paper.get("title", "")).strip()
    abstract = str(paper.get("abstract", "")).strip()
    blocks = [f"# {title}" if title else "", "## Abstract", abstract]
    for section in sections:
        section_name = str(section.get("section_name", "")).strip()
        if section_name:
            blocks.append(f"## {section_name}")
        paragraphs = section.get("paragraphs", [])
        if isinstance(paragraphs, Sequence) and not isinstance(paragraphs, (str, bytes)):
            blocks.extend(str(paragraph) for paragraph in paragraphs)
    return normalize_text("\n\n".join(block for block in blocks if block.strip()), "conservative")


def _answers_and_evidence(
    qa: Mapping[str, Any], document: str, paper_id: str
) -> tuple[list[str], list[dict[str, object]], str, bool]:
    accepted: list[str] = []
    evidence: list[dict[str, object]] = []
    answer_type = "unanswerable"
    answerable = False
    for annotation in _records(qa.get("answers", [])):
        raw_answer = annotation.get("answer", annotation)
        for answer in _records(raw_answer) or (
            [dict(raw_answer)] if isinstance(raw_answer, Mapping) else []
        ):
            if bool(answer.get("unanswerable")):
                continue
            value, current_type = _answer_value(answer)
            if value:
                accepted.append(value)
                answer_type = current_type
                answerable = True
            raw_evidence = answer.get("highlighted_evidence") or answer.get("evidence") or []
            if isinstance(raw_evidence, Sequence) and not isinstance(raw_evidence, (str, bytes)):
                for evidence_text in raw_evidence:
                    normalized = normalize_text(str(evidence_text), "conservative")
                    if not normalized or normalized.startswith("FLOAT SELECTED"):
                        continue
                    start = document.find(normalized)
                    if start < 0:
                        raise ValueError(
                            "QASPER evidence offset validation failed for "
                            f"{paper_id}: {normalized[:80]}"
                        )
                    item = {
                        "id": f"{paper_id}:{start}:{start + len(normalized)}",
                        "document_id": paper_id,
                        "start": start,
                        "end": start + len(normalized),
                    }
                    if item not in evidence:
                        evidence.append(item)
    return list(dict.fromkeys(accepted)), evidence, answer_type, answerable


def _answer_value(answer: Mapping[str, Any]) -> tuple[str, str]:
    free_form = str(answer.get("free_form_answer", "")).strip()
    if free_form:
        return free_form, "free_form"
    spans = answer.get("extractive_spans", [])
    if isinstance(spans, Sequence) and not isinstance(spans, (str, bytes)):
        joined = " ".join(str(value).strip() for value in spans if str(value).strip())
        if joined:
            return joined, "extractive"
    if "yes_no" in answer:
        return ("Yes" if bool(answer["yes_no"]) else "No"), "yes_no"
    return "", "unanswerable"


def _question_unanswerable(qa: Mapping[str, Any]) -> bool:
    for annotation in _records(qa.get("answers", [])):
        raw_answer = annotation.get("answer", annotation)
        answers = _records(raw_answer) or (
            [dict(raw_answer)] if isinstance(raw_answer, Mapping) else []
        )
        if any(bool(answer.get("unanswerable")) for answer in answers):
            return True
    return False


def _records(value: object) -> list[dict[str, Any]]:
    if isinstance(value, Sequence) and not isinstance(value, (str, bytes)):
        return [dict(item) for item in value if isinstance(item, Mapping)]
    if not isinstance(value, Mapping):
        return []
    if any(
        not isinstance(item, Sequence) or isinstance(item, (str, bytes)) for item in value.values()
    ):
        return [dict(value)]
    lengths = [
        len(item)
        for item in value.values()
        if isinstance(item, Sequence) and not isinstance(item, (str, bytes))
    ]
    if not lengths:
        return [dict(value)]
    count = min(lengths)
    return [
        {
            key: item[index]
            if isinstance(item, Sequence) and not isinstance(item, (str, bytes))
            else item
            for key, item in value.items()
        }
        for index in range(count)
    ]


if __name__ == "__main__":
    raise SystemExit(main())
